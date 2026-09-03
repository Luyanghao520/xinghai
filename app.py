# -*- coding: utf-8 -*-
"""星海艺术团官方网站（大系统）
- 招新子系统：复用原有 /api/signup、/admin、/cms、AI
- 成员信息子系统（员工式录入）：独立 members.db，学号为主键
- 统一登录 + 工作端：users.db，按角色
- 数据分离：各子系统独立库 / 文件，不混库
依赖 Flask（已预装）+ 标准库 sqlite3 / hashlib。
"""
import csv
import hashlib
import hmac
import io
import json
import os
import re
import shutil
import sqlite3
import threading
import time
from datetime import datetime
from html import escape

from flask import (Flask, request, jsonify, send_file, Response, redirect,
                     session, url_for)

BASE = os.path.dirname(os.path.abspath(__file__))
DB_REG = os.path.join(BASE, "registrations.db")   # 招新子系统
DB_MEM = os.path.join(BASE, "members.db")          # 成员信息子系统
DB_USR = os.path.join(BASE, "users.db")            # 统一账号
DB_REI = os.path.join(BASE, "reimburse.db")      # 报销子系统
DB_RES = os.path.join(BASE, "reserve.db")       # 预约子系统
DB_BUL = os.path.join(BASE, "bulletins.db")     # 活动通报子系统
DB_APP = os.path.join(BASE, "apply.db")          # 注册申请子系统
DB_AST = os.path.join(BASE, "assets.db")         # 资产管理子系统
KB_FILE = os.path.join(BASE, "kb.json")
CONTENT = os.path.join(BASE, "content.json")
STATIC = os.path.join(BASE, "static")
os.makedirs(os.path.join(STATIC, "uploads"), exist_ok=True)

# ===== 密钥配置（优先级：环境变量 > config.json > 内置默认）=====
_cfg = {}
_cfg_path = os.path.join(BASE, "config.json")
if os.path.exists(_cfg_path):
    try:
        with open(_cfg_path, encoding="utf-8") as _f:
            _cfg = json.load(_f)
    except Exception:
        _cfg = {}

ADMIN_KEY = os.environ.get("XINGHAI_ADMIN_KEY") or _cfg.get("ADMIN_KEY")  # 招新后台 / CMS 密码
SECRET = os.environ.get("XINGHAI_SECRET") or _cfg.get("SECRET")  # Flask session 密钥
if not ADMIN_KEY or not SECRET:
    raise RuntimeError(
        "XINGHAI_ADMIN_KEY 和 XINGHAI_SECRET 必须配置（环境变量或 config.json）"
    )

app = Flask(__name__, static_folder=STATIC, static_url_path="/static")
app.secret_key = SECRET

# Session cookie 安全：开发环境允许 HTTP，生产强制 HTTPS
SECURE_COOKIE = os.environ.get("FLASK_ENV") != "development"
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SECURE=SECURE_COOKIE,
    SESSION_COOKIE_SAMESITE="Lax",
    MAX_CONTENT_LENGTH=10 * 1024 * 1024,  # 10 MB 上传上限
)

# ===== 响应头：分级缓存 + 安全四件套（施工总案 A1）=====
@app.after_request
def xh_headers(resp):
    p = request.path
    if p.startswith("/assets/"):
        # 构建产物文件名带 hash，可永久缓存
        resp.headers["Cache-Control"] = "public, max-age=31536000, immutable"
    elif p.startswith("/static/") or p.startswith("/showcase/") or p.startswith("/u2026"):
        resp.headers["Cache-Control"] = "public, max-age=604800"  # 7 天
    else:
        resp.headers["Cache-Control"] = "no-cache"  # HTML 保持实时
    resp.headers.setdefault("X-Content-Type-Options", "nosniff")
    resp.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
    resp.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    resp.headers.setdefault("Permissions-Policy", "geolocation=(), microphone=(), camera=()")
    return resp

# ===== 官网内容编辑子系统（施工总案 A6 · site.db 独立库 + 二次口令）=====
SITE_EDIT_KEY = os.environ.get("XINGHAI_SITE_EDIT_KEY") or _cfg.get("SITE_EDIT_KEY")
# 优雅降级：未配置时不崩溃，相关接口返回提示（部署后 owner 在 config.json 补 SITE_EDIT_KEY 即启用）
DB_SITE = os.path.join(BASE, "site.db")
SITE_BACKUP_DIR = os.path.join(BASE, "data", "content_backups")
SITE_EDIT_TTL = 30 * 60      # 解锁有效期 30 分钟
SITE_MAX_FAIL = 5            # 连续失败上限
SITE_LOCK_SEC = 15 * 60      # 锁定 15 分钟
SITE_BACKUP_KEEP = 30        # 保留最近 30 份备份
_site_lock = threading.Lock()
_pub_cache = {"ts": 0.0, "data": None}

def site_audit(action, target="", detail=""):
    try:
        cx = db(DB_SITE)
        cx.execute(
            "INSERT INTO site_audit(ts,uid,name,ip,action,target,detail) VALUES(?,?,?,?,?,?,?)",
            (datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
             session.get("uid", ""), session.get("name", ""), _client_ip(),
             action, target, (detail or "")[:2000]))
        cx.commit(); cx.close()
    except Exception:
        pass

def _client_ip():
    return request.headers.get("X-Forwarded-For", request.remote_addr or "").split(",")[0].strip()

def _brief(v, n=120):
    s = v if isinstance(v, str) else json.dumps(v, ensure_ascii=False)
    return s[:n] + ("..." if len(s) > n else "")

def _guard_key():
    return "g:%s:%s" % (session.get("uid", "anon"), _client_ip())

def _guard_state():
    cx = db(DB_SITE)
    row = cx.execute("SELECT fails, locked_until FROM site_guard WHERE k=?", (_guard_key(),)).fetchone()
    cx.close()
    if not row:
        return 0.0, 0
    return float(row["locked_until"] or 0), int(row["fails"] or 0)

def _guard_fail():
    k = _guard_key(); now = time.time()
    cx = db(DB_SITE)
    row = cx.execute("SELECT fails FROM site_guard WHERE k=?", (k,)).fetchone()
    n = 1 if not row else int(row["fails"] or 0) + 1
    if not row:
        cx.execute("INSERT INTO site_guard(k,fails,first_fail,locked_until) VALUES(?,?,?,?)", (k, n, now, 0))
    else:
        cx.execute("UPDATE site_guard SET fails=? WHERE k=?", (n, k))
    if n >= SITE_MAX_FAIL:
        cx.execute("UPDATE site_guard SET locked_until=?, fails=0 WHERE k=?", (now + SITE_LOCK_SEC, k))
        cx.commit(); cx.close()
        site_audit("LOCK", "", "连续 %d 次口令错误，锁定 %d 分钟" % (n, SITE_LOCK_SEC // 60))
        return n, True
    cx.commit(); cx.close()
    site_audit("UNLOCK_FAIL", "", "第 %d 次失败" % n)
    return n, False

def _guard_reset():
    cx = db(DB_SITE)
    cx.execute("DELETE FROM site_guard WHERE k=?", (_guard_key(),))
    cx.commit(); cx.close()

def _site_unlocked():
    uid = session.get("uid")
    if not uid:
        return False
    if session.get("site_edit_uid") != uid:
        return False
    return float(session.get("site_edit_until", 0) or 0) > time.time()

def site_edit_required(f):
    """第三层鉴权：二次口令（须与 @login_required 串联，@app.route 最上）"""
    from functools import wraps
    @wraps(f)
    def w(*a, **k):
        if "uid" not in session:
            return jsonify({"ok": False, "error": "未登录"}), 401
        if not SITE_EDIT_KEY:
            return jsonify({"ok": False, "error": "编辑口令未配置（config.json 缺 SITE_EDIT_KEY）"}), 503
        if not _site_unlocked():
            return jsonify({"ok": False, "error": "需要二次验证",
                            "code": "STEP_UP_REQUIRED"}), 403
        session["site_edit_until"] = time.time() + SITE_EDIT_TTL  # 滑动续期
        return f(*a, **k)
    return w

def _backup_content(path):
    if not os.path.exists(path):
        return None
    os.makedirs(SITE_BACKUP_DIR, exist_ok=True)
    prefix = os.path.basename(path).replace(".json", "") + "-"
    ts = datetime.now().strftime("%Y%m%d-%H%M%S-%f")  # 微秒级，防同秒互覆
    dst = os.path.join(SITE_BACKUP_DIR, prefix + ts + ".json")
    shutil.copy2(path, dst)
    try:
        fs = sorted([f for f in os.listdir(SITE_BACKUP_DIR) if f.startswith(prefix)], reverse=True)
        for old in fs[SITE_BACKUP_KEEP:]:
            os.remove(os.path.join(SITE_BACKUP_DIR, old))
    except Exception:
        pass
    return dst

def save_content_atomic(path, obj):
    """写临时文件 + os.replace 原子替换，避免写一半崩溃导致 JSON 损坏"""
    with _site_lock:
        _backup_content(path)
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(obj, f, ensure_ascii=False, indent=2)
            f.flush()
            os.fsync(f.fileno())
        os.replace(tmp, path)
        _pub_cache["ts"] = 0.0  # 清运行时缓存

# ============ 工具 ============
def db(path):
    """SQLite 连接（施工总案 A2：WAL + busy_timeout，招新高峰并发写不炸）"""
    cx = sqlite3.connect(path, timeout=15)
    cx.row_factory = sqlite3.Row
    cx.execute("PRAGMA journal_mode=WAL")
    cx.execute("PRAGMA busy_timeout=15000")
    cx.execute("PRAGMA synchronous=NORMAL")
    return cx

def hx(s):  # 学号哈希（加盐）
    return hashlib.sha256((SECRET + str(s)).encode("utf-8")).hexdigest()

def _session_pv_valid():
    """会话内的密码版本号与库中一致（改密后旧会话全部失效）"""
    cx = db(DB_USR)
    row = cx.execute("SELECT pv FROM users WHERE xh=?", (session["uid"],)).fetchone()
    cx.close()
    cur_pv = (row["pv"] or 0) if row else 0
    return session.get("pv", 0) == cur_pv

def login_required(f):
    from functools import wraps
    @wraps(f)
    def w(*a, **k):
        if "uid" not in session:
            if request.path.startswith("/api"):
                return jsonify({"ok": False, "error": "未登录"}), 401
            return redirect("/login")
        if not _session_pv_valid():
            session.clear()
            if request.path.startswith("/api"):
                return jsonify({"ok": False, "error": "会话已失效，请重新登录"}), 401
            return redirect("/login")
        return f(*a, **k)
    return w

def chair_required(f):
    from functools import wraps
    @wraps(f)
    def w(*a, **k):
        if session.get("role") not in ("主席", "副主席"):
            return jsonify({"ok": False, "error": "仅主席/副主席可操作"}), 403
        return f(*a, **k)
    return w

# 图片内容嗅探（施工总案 A4：零依赖 magic-bytes，拒绝改名伪装的假图片）
_IMG_MAGIC = (
    ("jpeg", b"\xff\xd8\xff"),
    ("png",  b"\x89PNG\r\n\x1a\n"),
    ("gif",  b"GIF87a"),
    ("gif",  b"GIF89a"),
)
def sniff_image(raw):
    if raw[:4] == b"RIFF" and raw[8:12] == b"WEBP":
        return "webp"
    for fmt, m in _IMG_MAGIC:
        if raw.startswith(m):
            return fmt
    return None


# ============ 初始化 ============
def init_reg():
    cx = db(DB_REG)
    cx.execute("""CREATE TABLE IF NOT EXISTS registrations(
        id INTEGER PRIMARY KEY AUTOINCREMENT, time TEXT, target TEXT, name TEXT,
        gender TEXT, birth TEXT, campus TEXT, college TEXT, major TEXT,
        phone TEXT, wechat TEXT, email TEXT, skill TEXT, motive TEXT, adjust INTEGER)""")
    cx.commit()
    jf = os.path.join(BASE, "registrations.json")
    if os.path.exists(jf) and os.path.getsize(jf) > 2:
        try:
            rows = json.load(open(jf, encoding="utf-8"))
            cur = cx.execute("SELECT COUNT(*) AS c FROM registrations").fetchone()["c"]
            if cur == 0:
                for r in rows:
                    cx.execute("INSERT INTO registrations (time,target,name,gender,birth,campus,college,major,phone,wechat,email,skill,motive,adjust) "
                                 "VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
                                 (r.get("time"), r.get("name"), r.get("gender"),
                                  r.get("birth"), r.get("campus"), r.get("college"), r.get("major"),
                                  r.get("phone"), r.get("wechat"), r.get("email"),
                                  r.get("skill"), r.get("motive"), r.get("adjust"), 1 if r.get("adjust") else 0))
                cx.commit()
            os.rename(jf, jf + ".migrated")
        except Exception as e:
            app.logger.warning("migrate json failed: %s", e)
    cx.close()

MEM_COLS = ["xh", "name", "gender", "campus", "college", "major",
             "phone", "wechat", "email", "role", "dept",
             "join_date", "grade", "status", "skill", "note", "position"]
MEM_LABELS = {"xh": "学号", "name": "姓名", "gender": "性别", "campus": "校区",
              "college": "学院", "major": "专业/班级", "phone": "手机号", "wechat": "微信",
              "email": "邮箱", "role": "角色", "dept": "所属部门/团队", "join_date": "入团时间",
              "grade": "届别", "status": "状态", "skill": "特长/经历", "note": "备注", "position": "职务"}
POSITIONS = ["教师","主席","副主席","团秘","部长","副部长","组长","团长","副团长","团干","成员"]

def init_mem():
    cx = db(DB_MEM)
    cx.execute("""CREATE TABLE IF NOT EXISTS members(
        xh TEXT PRIMARY KEY, name TEXT, gender TEXT, campus TEXT, college TEXT,
        major TEXT, phone TEXT, wechat TEXT, email TEXT, role TEXT, dept TEXT,
        join_date TEXT, grade TEXT, status TEXT, skill TEXT, note TEXT,
        position TEXT DEFAULT '成员', updated TEXT)""")
    try: cx.execute("ALTER TABLE members ADD COLUMN position TEXT DEFAULT '成员'")
    except sqlite3.OperationalError: pass  # 列已存在则跳过
    cx.commit(); cx.close()

def init_usr():
    cx = db(DB_USR)
    cx.execute("""CREATE TABLE IF NOT EXISTS users(
        xh TEXT PRIMARY KEY, name TEXT, role TEXT, pwd TEXT, campus TEXT, status TEXT)""")
    try: cx.execute("ALTER TABLE users ADD COLUMN pv INTEGER DEFAULT 0")  # 密码版本号
    except sqlite3.OperationalError: pass  # 列已存在则跳过
    cx.commit()
    if cx.execute("SELECT COUNT(*) AS c FROM users").fetchone()["c"] == 0:
        import secrets as _secrets
        _tmp = _secrets.token_urlsafe(12)  # 施工总案 A7：初始密码随机化，消除硬编码后门
        cx.execute("INSERT INTO users (xh,name,role,pwd,campus,status) VALUES (?,?,?,?,?,?)",
                   ("000000000", "系统管理员", "主席", hx(_tmp), "", "在团"))
        cx.commit()
        print("[INIT] 已创建初始管理员 000000000，临时密码：%s  ⚠️ 请立即登录后修改" % _tmp)
    cx.close()

# ============ 报销子系统（员工式，独立库） ============
def init_rei():
    cx = db(DB_REI)
    cx.execute("""CREATE TABLE IF NOT EXISTS rei(
        id INTEGER PRIMARY KEY AUTOINCREMENT, xh TEXT, name TEXT, title TEXT,
        category TEXT, amount REAL, note TEXT, status TEXT,
        created TEXT, updated TEXT)""")
    cx.commit(); cx.close()

REI_COLS = ["xh", "name", "title", "category", "amount", "note", "status"]

# ============ 预约子系统（员工式，独立库） ============
def init_res():
    cx = db(DB_RES)
    cx.execute("""CREATE TABLE IF NOT EXISTS resv(
        id INTEGER PRIMARY KEY AUTOINCREMENT, xh TEXT, name TEXT, item TEXT,
        type TEXT, start_time TEXT, end_time TEXT, purpose TEXT,
        status TEXT, created TEXT, updated TEXT)""")
    cx.commit(); cx.close()

RES_COLS = ["xh", "name", "item", "type", "start_time", "end_time", "purpose", "status"]

# ============ 活动通报子系统（独立库） ============
def init_bul():
    cx = db(DB_BUL)
    cx.execute("""CREATE TABLE IF NOT EXISTS bul(
        id INTEGER PRIMARY KEY AUTOINCREMENT, author TEXT, author_xh TEXT,
        title TEXT, body TEXT, level TEXT, pinned INTEGER, created TEXT)""")
    cx.commit(); cx.close()


# ============ 注册申请子系统（公开提交 + 管理员审批）============
def init_apply():
    cx = db(DB_APP)
    cx.execute("""CREATE TABLE IF NOT EXISTS applies(
        id INTEGER PRIMARY KEY AUTOINCREMENT, xh TEXT, pwd TEXT, name TEXT,
        campus TEXT, status TEXT DEFAULT '待审', created TEXT, updated TEXT)""")
    cx.commit(); cx.close()
init_apply()

# ============ 资产管理子系统（独立库 assets.db：资产 + 借还 + 采购审批）============
def init_ast():
    cx = db(DB_AST)
    cx.execute("""CREATE TABLE IF NOT EXISTS assets(
        id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, category TEXT,
        qty INTEGER DEFAULT 1, location TEXT, person TEXT, status TEXT DEFAULT '在库',
        created TEXT, updated TEXT)""")
    cx.execute("""CREATE TABLE IF NOT EXISTS borrows(
        id INTEGER PRIMARY KEY AUTOINCREMENT, asset_id INTEGER, asset_name TEXT,
        xh TEXT, name TEXT, qty INTEGER DEFAULT 1, note TEXT,
        borrow_date TEXT, return_date TEXT, status TEXT DEFAULT '借出', created TEXT)""")
    cx.execute("""CREATE TABLE IF NOT EXISTS procurements(
        id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, category TEXT,
        qty INTEGER DEFAULT 1, reason TEXT, xh TEXT, pname TEXT,
        status TEXT DEFAULT '待审', created TEXT, updated TEXT)""")
    cx.commit(); cx.close()
init_ast()

# ============ 离团存档表（members.db）============
def init_alumni():
    cx = db(DB_MEM)
    cx.execute("""CREATE TABLE IF NOT EXISTS alumni(
        xh TEXT, name TEXT, gender TEXT, campus TEXT, college TEXT, major TEXT,
        phone TEXT, wechat TEXT, email TEXT, role TEXT, dept TEXT,
        join_date TEXT, grade TEXT, position TEXT, leave_date TEXT, note TEXT)""")
    cx.commit(); cx.close()
init_alumni()

# ============ 招新淘汰封存表（registrations.db）============
def init_archive_regs():
    cx = db(DB_REG)
    cx.execute("""CREATE TABLE IF NOT EXISTS archive_regs(
        id INTEGER PRIMARY KEY, time TEXT, target TEXT, name TEXT, gender TEXT,
        birth TEXT, campus TEXT, college TEXT, major TEXT,
        phone TEXT, wechat TEXT, email TEXT, skill TEXT, motive TEXT, adjust INTEGER,
        status TEXT, archived_at TEXT)""")
    cx.commit(); cx.close()
init_archive_regs()

# ============ 小星知识库（AI 答疑素材，可后台填充）============
DEFAULT_KB = [
    {"keywords":["你好","您好","在吗"],"answer":"你好呀～我是星海艺术团吉祥物 <b>小星</b> ✨"},
    {"keywords":["介绍","星海","艺术团","组织"],"answer":"星海艺术团是上海立信会计金融学院校级艺术团体，涵盖声乐、器乐、舞蹈、主持、礼仪、话剧、书法等多个方向。"},
    {"keywords":["招新","报名","加入","怎么进"],"answer":"每年 9 月开学后统一招新，关注 QQ 招新群或公众号通知，或点击 <a href='/register'>注册页面</a> 提交申请。"},
    {"keywords":["部门","方向","声乐","器乐","舞蹈","主持","礼仪","话剧","书法","合唱","乐团"],"answer":"星海艺术团下设声乐团、合唱团、民乐团、管弦乐团、舞蹈团、主持团、礼仪团、话剧团、书法社等方向，具体可查看招新简章。"},
    {"keywords":["排练","时间","地点"],"answer":"各团队排练时间不同，一般在工作日晚间或周末，地点在文翔路校区/上川路校区排练厅，具体由各团长通知。"},
    {"keywords":["演出","活动","比赛"],"answer":"星海艺术团定期参加校内外演出，如迎新晚会、五四汇演、上海市大学生艺术展演等，是展示才艺的绝佳平台。"},
    {"keywords":["加分","学分","综测","二课","素拓"],"answer":"参与艺术团活动可按学校规定获得二课学分、社会实践证明及综测加分，具体以当年学生处政策为准。"},
    {"keywords":["指导老师","教师","主席","团长"],"answer":"艺术团由校团委指导，各团队均有专业教师指导；主席团由团长、副团长等组成，负责日常管理。"},
    {"keywords":["请假","旷训","签到"],"answer":"排练请假需提前向团长报备并说明原因，无故缺席将记录考勤影响综测；请假请联系团长并保留记录。"},
    {"keywords":["面试","考核","选拔"],"answer":"招新报名后由各团组织面试/考核，内容因方向而异：声乐需清唱一段，器乐需演奏一首，舞蹈需展示一段，主持需即兴主持等。"},
    {"keywords":["校区","文翔路","上川路","浦东","松江"],"answer":"星海艺术团在文翔路校区（松江）和上川路校区（浦东）均设有团队，部分团队跨校区活动。"},
    {"keywords":["办公室","值班","地点"],"answer":"团办公室位于排练厅旁，值班时间由主席团安排，具体可咨询各团部长或查看公告栏。"},
    {"keywords":["服装","道具","演出服"],"answer":"演出服装和道具由艺术团统一管理，演出前由负责人到办公室统一领取，用后需按时归还。"},
    {"keywords":["换届","任职","竞选"],"answer":"艺术团每年 6 月进行换届选举，岗位包括团长、副团长、部长、副部长等，需提交申请并参加竞聘答辩。"},
    {"keywords":["报销","经费","发票","材料"],"answer":"演出道具、服装等需报销的请通过工作端「报销」模块提交，附发票照片及相关说明，由主席团审批。"},
    {"keywords":["通知","消息","公告","群","QQ","微信"],"answer":"重要通知通过 QQ 通知群和微信群发布，请确保已加入相应群组并关注群公告。"},
    {"keywords":["退团","离团","退出"],"answer":"如需离团，请先与团长/主席团沟通，之后由管理员在工作端操作离团登记。"},
    {"keywords":["证明","证书","聘书","荣誉"],"answer":"学年结束时统一发放文艺骨干证明、聘书等，具体时间和领取方式届时通知。"},
    {"keywords":["乐器","设备","资产","音响","灯光"],"answer":"艺术团拥有钢琴、音响、灯光等专业设备，借用需通过工作端「资产管理」提交申请，按时归还。"},
    {"keywords":["场地","排练厅","琴房","借场地"],"answer":"排练厅和琴房可通过工作端「场地预约」模块申请使用，需提前至少1天提交申请，审批通过后使用。"},
    {"keywords":["培训","学习","大师课","讲座"],"answer":"艺术团不定期邀请专业老师开设大师课和培训讲座，所有团员均可参加，具体安排见通知。"},
    {"keywords":["报销流程","怎么报销","报销步骤"],"answer":"报销流程：1)登录工作端 2)进入报销页面 3)填写报销单并上传发票照片 4)提交等待主席团审批 5)审批通过后打印报销单。"},
    {"keywords":["联系方式","电话","邮箱","QQ群"],"answer":"艺术团官方QQ群：请关注招新公告获取最新群号；办公室主任邮箱：可在工作端成员列表查看联系方式。"},
    {"keywords":["注册","账号","登录","忘记密码"],"answer":"首次使用请通过「注册」页面提交学号等信息，审核通过后即可登录；忘记密码请联系主席团重置。"}
]
def init_kb():
    if os.path.exists(KB_FILE): return
    data = [{"id":i+1,"keywords":it["keywords"],"answer":it["answer"]} for i,it in enumerate(DEFAULT_KB)]
    with open(KB_FILE,"w",encoding="utf-8") as f: json.dump(data,f,ensure_ascii=False,indent=2)
init_kb()

# ============ 招新子系统（复用） ============
@app.route("/recruit")
def recruit():
    return send_file(os.path.join(BASE, "recruit.html"))

@app.route("/api/signup", methods=["POST"])
def signup():
    try:
        d = request.get_json(force=True, silent=True) or {}
    except Exception:
        return jsonify({"ok": False, "error": "数据格式错误"}), 400
    name = (d.get("name") or "").strip()
    campus = (d.get("campus") or "").strip()
    gender = (d.get("gender") or "").strip()
    college = (d.get("college") or "").strip()
    major = (d.get("major") or "").strip()
    phone = (d.get("phone") or "").strip()
    if not (name and campus and gender and college and major and phone):
        return jsonify({"ok": False, "error": "必填项不完整"}), 400
    if not phone.isdigit() or len(phone) != 11:
        return jsonify({"ok": False, "error": "手机号格式不正确"}), 400
    cx = db(DB_REG)
    cols = ["time","target","name","gender","birth","campus","college","major","phone","wechat","email","skill","motive","adjust"]
    vals = [
        datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        (d.get("target") or "").strip(), name, gender, (d.get("birth") or "").strip(),
        campus, college, major, phone,
        (d.get("wechat") or "").strip(), (d.get("email") or "").strip(),
        (d.get("skill") or "").strip(), (d.get("motive") or "").strip(),
        1 if d.get("adjust") else 0,
    ]
    assert len(cols) == len(vals), (len(cols), len(vals))
    cx.execute("INSERT INTO registrations (" + ",".join(cols) + ") VALUES (" + ",".join("?" * len(cols)) + ")", vals)
    cx.commit()
    total = cx.execute("SELECT COUNT(*) AS c FROM registrations").fetchone()["c"]
    cx.close()
    export_csv()
    return jsonify({"ok": True, "count": total})

@app.route("/api/delete", methods=["POST"])
@login_required
@chair_required
def delete():
    rid = (request.get_json(silent=True) or {}).get("id") or request.args.get("id")
    if not rid:
        return jsonify({"ok": False, "error": "缺少 id"}), 400
    cx = db(DB_REG); cx.execute("DELETE FROM registrations WHERE id=?", (rid,)); cx.commit(); cx.close()
    export_csv()
    return jsonify({"ok": True})

def all_rows(q=None):
    cx = db(DB_REG)
    if q:
        like = "%" + q + "%"
        rows = cx.execute("SELECT * FROM registrations WHERE name LIKE ? OR phone LIKE ? OR target LIKE ? "
                       "OR college LIKE ? OR major LIKE ? ORDER BY id DESC", (like, like, like, like, like)).fetchall()
    else:
        rows = cx.execute("SELECT * FROM registrations ORDER BY id DESC").fetchall()
    cx.close()
    return [dict(r) for r in rows]

_LAST_CSV = 0.0

def export_csv(force=False):
    """节流：报名/删除触发的调用 30s 内最多写一次；导出给用户看时传 force=True"""
    global _LAST_CSV
    import time as _t
    if not force and _t.time() - _LAST_CSV < 30:
        return
    _LAST_CSV = _t.time()
    rows = all_rows()
    path = os.path.join(BASE, "registrations.csv")
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=["time", "target", "name", "gender", "birth", "campus",
                                            "college", "major", "phone", "wechat", "email", "skill", "motive", "adjust"])
        w.writeheader()
        for r in rows:
            w.writerow({k: r.get(k, "") for k in w.fieldnames})

@app.route("/admin")
@login_required
@chair_required
def admin():
    if request.args.get("export") == "csv":
        export_csv(force=True)  # 导出给用户看，必须全量最新
        return send_file(os.path.join(BASE, "registrations.csv"), mimetype="text/csv",
                         as_attachment=True, download_name="星海艺术团报名.csv")
    return send_file(os.path.join(BASE, "admin.html"))

@app.route("/api/admin/rows")
@login_required
@chair_required
def api_admin_rows():
    """招新后台数据接口（会话 + 主席团鉴权）：统计 + 明细一次拉取"""
    q = (request.args.get("q") or "").strip()
    rows = all_rows(q if q else None)
    cx = db(DB_REG)
    total = cx.execute("SELECT COUNT(*) AS c FROM registrations").fetchone()["c"]
    by_campus = dict((r["campus"] or "未填", r["c"]) for r in
                      cx.execute("SELECT campus, COUNT(*) AS c FROM registrations GROUP BY campus"))
    by_target = dict((r["target"] or "未填", r["c"]) for r in
                       cx.execute("SELECT target, COUNT(*) AS c FROM registrations GROUP BY target"))
    cx.close()
    return jsonify({"ok": True, "total": total, "by_campus": by_campus,
                    "by_target": by_target, "rows": rows})

# ============ 内容库（CMS，复用） ============
def load_content():
    try:
        if os.path.exists(CONTENT):
            with open(CONTENT, encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return {}

def save_content(obj):
    with open(CONTENT, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)

@app.route("/api/content")
def api_content():
    return jsonify(load_content())

@app.route("/api/content", methods=["POST"])
def api_content_set():
    if request.args.get("key") != ADMIN_KEY and (request.get_json(silent=True) or {}).get("key") != ADMIN_KEY:
        return jsonify({"ok": False, "error": "无权限"}), 403
    data = request.get_json(force=True, silent=True) or {}
    cur = load_content(); cur.update(data); save_content(cur)
    return jsonify({"ok": True})

@app.route("/api/upload", methods=["POST"])
def api_upload():
    if request.args.get("key") != ADMIN_KEY:
        return jsonify({"ok": False, "error": "无权限"}), 403
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "未收到文件"}), 400
    name = (f.filename or "img.png").rsplit(".", 1)
    ext = (name[-1] if len(name) > 1 else "png").lower()
    if ext not in ("png", "jpg", "jpeg", "gif", "webp"):
        return jsonify({"ok": False, "error": "仅支持图片文件"}), 400
    raw = f.read()
    if len(raw) > 5 * 1024 * 1024:
        return jsonify({"ok": False, "error": "图片不能超过 5MB"}), 400
    if not sniff_image(raw):
        return jsonify({"ok": False, "error": "文件内容不是有效图片"}), 400
    f.stream = io.BytesIO(raw)  # sniff 消耗了流，重置后再 save
    safe = "u" + datetime.now().strftime("%Y%m%d%H%M%S") + "_" + os.urandom(4).hex() + "." + ext
    f.save(os.path.join(STATIC, "uploads", safe))
    return jsonify({"ok": True, "url": "/static/uploads/" + safe})

@app.route("/cms")
def cms():
    return send_file(os.path.join(BASE, "cms.html"))

# ============ 统一登录 ============
@app.route("/login", methods=["GET"])
def login_page():
    if "uid" in session:
        return redirect("/work")
    return send_file(os.path.join(BASE, "login.html"))

@app.route("/api/login", methods=["POST"])
def api_login():
    d = request.get_json(force=True, silent=True) or {}
    xh = (d.get("xh") or "").strip()
    pwd = (d.get("pwd") or "").strip()
    if not (xh and pwd):
        return jsonify({"ok": False, "error": "请输入学号/工号和密码"}), 400
    cx = db(DB_USR)
    u = cx.execute("SELECT * FROM users WHERE xh=?", (xh,)).fetchone()
    cx.close()
    if not u or u["pwd"] != hx(pwd):
        return jsonify({"ok": False, "error": "学号/工号或密码错误"}), 401
    session["uid"] = u["xh"]; session["name"] = u["name"]; session["role"] = u["role"]
    session["pv"] = u["pv"] or 0
    return jsonify({"ok": True, "name": u["name"], "role": u["role"]})

@app.route("/api/reset", methods=["POST"])
def api_reset():
    """密码重置：旧密码验证 + 新密码规则校验；pv+1 使所有设备的旧会话全部失效"""
    d = request.get_json(force=True, silent=True) or {}
    xh = (d.get("xh") or "").strip()
    old = d.get("old_pwd") or ""
    new = d.get("new_pwd") or ""
    if not (xh and old and new):
        return jsonify({"ok": False, "error": "请填写完整"}), 400
    if len(new) < 8 or not any(c.islower() for c in new) \
            or not any(c.isupper() for c in new) or not any(c.isdigit() for c in new):
        return jsonify({"ok": False, "error": "新密码须≥8位，包含大小写字母和数字"}), 400
    if new == old:
        return jsonify({"ok": False, "error": "新密码不能与旧密码相同"}), 400
    cx = db(DB_USR)
    u = cx.execute("SELECT * FROM users WHERE xh=?", (xh,)).fetchone()
    if not u or u["pwd"] != hx(old):
        cx.close()
        return jsonify({"ok": False, "error": "学号或旧密码错误"}), 401
    cx.execute("UPDATE users SET pwd=?, pv=COALESCE(pv,0)+1 WHERE xh=?", (hx(new), xh))
    cx.commit(); cx.close()
    session.clear()  # 本设备的会话一并结束，需用新密码重新登录
    return jsonify({"ok": True, "msg": "密码已更新，所有设备的会话已结束"})

@app.route("/reset", methods=["GET"])
def reset_page():
    return send_file(os.path.join(BASE, "reset.html"))

@app.route("/api/me")
def api_me():
    if "uid" not in session:
        return jsonify({"ok": False})
    if not _session_pv_valid():
        session.clear()
        return jsonify({"ok": False, "expired": True})
    return jsonify({"ok": True, "name": session.get("name"),
                    "role": session.get("role"), "uid": session.get("uid")})

@app.route("/logout", methods=["GET", "POST"])  # POST 优先（防 <img src> CSRF），GET 兼容现有前端链接
def logout():
    session.clear()
    return redirect("/login")

# ============ 成员信息子系统（员工式录入） ============
@app.route("/work")
@login_required
def work():
    return send_file(os.path.join(BASE, "work.html"))

@app.route("/work/members")
@login_required
def members_page():
    return send_file(os.path.join(BASE, "members.html"))

@app.route("/api/members")
@login_required
def api_members():
    q = (request.args.get("q") or "").strip()
    campus = (request.args.get("campus") or "").strip()
    role = (request.args.get("role") or "").strip()
    cx = db(DB_MEM)
    sql = "SELECT * FROM members WHERE 1=1"
    args = []
    if q:
        sql += " AND (xh LIKE ? OR name LIKE ? OR dept LIKE ? OR college LIKE ? OR major LIKE ?)"
        like = "%" + q + "%"
        args += [like, like, like, like, like]
    if campus:
        sql += " AND campus=?"; args.append(campus)
    if role:
        sql += " AND role=?"; args.append(role)
    sql += " ORDER BY xh"
    rows = [dict(r) for r in cx.execute(sql, args).fetchall()]
    cx.close()
    return jsonify(rows)

@app.route("/api/members", methods=["POST"])
@login_required
def api_members_save():
    d = request.get_json(force=True, silent=True) or {}
    xh = (d.get("xh") or "").strip()
    if not (xh.isdigit() and len(xh) == 9):
        return jsonify({"ok": False, "error": "学号须为 9 位数字（如 251400143）"}), 400
    vals = {c: (d.get(c) or "").strip() for c in MEM_COLS if c != "xh"}
    vals["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx = db(DB_MEM)
    exists = cx.execute("SELECT 1 FROM members WHERE xh=?", (xh,)).fetchone()
    if exists:
        cx.execute("UPDATE members SET " + ",".join("%s=?" % c for c in MEM_COLS if c != "xh") +
                   ",updated=? WHERE xh=?", tuple(vals[c] for c in MEM_COLS if c != "xh") + (vals["updated"], xh))
    else:
        cols = ["xh"] + [c for c in MEM_COLS if c != "xh"] + ["updated"]
        cx.execute("INSERT INTO members (" + ",".join(cols) + ") VALUES (" + ",".join("?" * len(cols)) + ")",
                   (xh,) + tuple(vals[c] for c in MEM_COLS if c != "xh") + (vals["updated"],))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/members/<xh>", methods=["DELETE"])
@chair_required
def api_members_del(xh):
    cx = db(DB_MEM); cx.execute("DELETE FROM members WHERE xh=?", (xh,)); cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/members/import", methods=["POST"])
@chair_required
def api_members_import():
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "未收到文件"}), 400
    raw = f.read().decode("utf-8-sig", errors="ignore")
    reader = csv.DictReader(io.StringIO(raw))
    hmap = {"学号": "xh", "姓名": "name", "性别": "gender", "校区": "campus", "学院": "college",
             "专业/班级": "major", "专业班级": "major", "手机号": "phone", "微信": "wechat", "邮箱": "email",
             "角色": "role", "所属部门/团队": "dept", "所属部门团队": "dept", "入团时间": "join_date",
             "届别": "grade", "状态": "status", "特长/经历": "skill", "特长经历": "skill", "备注": "note"}
    n = 0
    cx = db(DB_MEM)
    for row in reader:
        m = {hmap.get(k, k): (v or "").strip() for k, v in row.items() if hmap.get(k, k) in MEM_COLS}
        xh = (m.get("xh") or "").strip()
        if not (xh.isdigit() and len(xh) == 9):
            continue
        vals = {c: m.get(c, "") for c in MEM_COLS if c != "xh"}
        vals["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        if cx.execute("SELECT 1 FROM members WHERE xh=?", (xh,)).fetchone():
            cx.execute("UPDATE members SET " + ",".join("%s=?" % c for c in MEM_COLS if c != "xh") +
                       ",updated=? WHERE xh=?", tuple(vals[c] for c in MEM_COLS if c != "xh") + (vals["updated"], xh))
        else:
            cols = ["xh"] + [c for c in MEM_COLS if c != "xh"] + ["updated"]
            cx.execute("INSERT INTO members (" + ",".join(cols) + ") VALUES (" + ",".join("?" * len(cols)) + ")",
                       (xh,) + tuple(vals[c] for c in MEM_COLS if c != "xh") + (vals["updated"],))
        n += 1
    cx.commit(); cx.close()
    return jsonify({"ok": True, "imported": n})

@app.route("/api/members/export")
@login_required
def api_members_export():
    cx = db(DB_MEM)
    rows = cx.execute("SELECT * FROM members ORDER BY xh").fetchall()
    cx.close()
    buf = io.StringIO()
    w = csv.DictWriter(buf, fieldnames=MEM_COLS)
    w.writeheader()
    for r in rows:
        w.writerow({c: r[c] for c in MEM_COLS})
    return Response(buf.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition": "attachment; filename=members.csv"})

# ============ 官网首页（学生端） ============
# =========== 报销子系统（员工式，独立库） ===========
@app.route("/work/reimburse")
@login_required
def reimburse_page():
    return send_file(os.path.join(BASE, "reimburse.html"))

@app.route("/api/reimburse")
@login_required
def api_reimburse():
    mine = request.args.get("mine")
    cx = db(DB_REI)
    if mine:
        rows = cx.execute("SELECT * FROM rei WHERE xh=? ORDER BY id DESC", (session["uid"],)).fetchall()
    else:
        rows = cx.execute("SELECT * FROM rei ORDER BY id DESC").fetchall()
    cx.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/reimburse", methods=["POST"])
@login_required
def api_reimburse_save():
    d = request.get_json(force=True, silent=True) or {}
    title = (d.get("title") or "").strip()
    category = (d.get("category") or "").strip()
    amt = d.get("amount")
    if amt is None or amt == "":
        return jsonify({"ok": False, "error": "金额须为数字"}), 400
    try:
        amt = float(amt)
    except Exception:
        return jsonify({"ok": False, "error": "金额须为数字"}), 400
    if not (title and category and amt > 0):
        return jsonify({"ok": False, "error": "请填写事项、类别与正数金额"}), 400
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx = db(DB_REI)
    cols = ["xh", "name", "title", "category", "amount", "note", "status", "created", "updated"]
    vals = [session["uid"], session["name"], title, category, amt,
             (d.get("note") or "").strip(), "待审", now, now]
    cx.execute("INSERT INTO rei (" + ",".join(cols) + ") VALUES (" + ",".join("?" * len(cols)) + ")", vals)
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/reimburse/<int:rid>", methods=["DELETE"])
@login_required
def api_reimburse_del(rid):
    cx = db(DB_REI)
    row = cx.execute("SELECT xh FROM rei WHERE id=?", (rid,)).fetchone()
    if row and (row["xh"] == session["uid"] or session.get("role") in ("主席", "副主席")):
        cx.execute("DELETE FROM rei WHERE id=?", (rid,)); cx.commit()
    cx.close()
    return jsonify({"ok": True})

@app.route("/api/reimburse/<int:rid>/status", methods=["POST"])
@login_required
def api_reimburse_status(rid):
    if session.get("role") not in ("主席", "副主席"):
        return jsonify({"ok": False, "error": "仅主席/副主席可审批"}), 403
    d = request.get_json(force=True, silent=True) or {}
    st = (d.get("status") or "").strip()
    if st not in ("待审", "已通过", "已驳回"):
        return jsonify({"ok": False, "error": "状态非法"}), 400
    cx = db(DB_REI)
    cx.execute("UPDATE rei SET status=?, updated=? WHERE id=?",
               (st, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), rid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

# =========== 预约子系统（员工式，独立库） ===========
@app.route("/work/reserve")
@login_required
def reserve_page():
    return send_file(os.path.join(BASE, "reserve.html"))

@app.route("/api/reserve")
@login_required
def api_reserve():
    mine = request.args.get("mine")
    cx = db(DB_RES)
    if mine:
        rows = cx.execute("SELECT * FROM resv WHERE xh=? ORDER BY id DESC", (session["uid"],)).fetchall()
    else:
        rows = cx.execute("SELECT * FROM resv ORDER BY id DESC").fetchall()
    cx.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/reserve", methods=["POST"])
@login_required
def api_reserve_save():
    d = request.get_json(force=True, silent=True) or {}
    item = (d.get("item") or "").strip()
    type_ = (d.get("type") or "").strip()
    start = (d.get("start_time") or "").strip()
    end = (d.get("end_time") or "").strip()
    if not (item and type_ and start and end):
        return jsonify({"ok": False, "error": "请填写对象、名称与起止时间"}), 400
    if start >= end:
        return jsonify({"ok": False, "error": "结束时间必须晚于开始时间"}), 400
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx = db(DB_RES)
    cols = ["xh", "name", "item", "type", "start_time", "end_time", "purpose", "status", "created", "updated"]
    vals = [session["uid"], session["name"], item, type_, start, end,
            (d.get("purpose") or "").strip(), "待确认", now, now]
    cx.execute("INSERT INTO resv (" + ",".join(cols) + ") VALUES (" + ",".join("?" * len(cols)) + ")", vals)
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/reserve/<int:rid>", methods=["DELETE"])
@login_required
def api_reserve_del(rid):
    cx = db(DB_RES)
    row = cx.execute("SELECT xh FROM resv WHERE id=?", (rid,)).fetchone()
    if row and (row["xh"] == session["uid"] or session.get("role") in ("主席", "副主席")):
        cx.execute("DELETE FROM resv WHERE id=?", (rid,)); cx.commit()
    cx.close()
    return jsonify({"ok": True})

@app.route("/api/reserve/<int:rid>/status", methods=["POST"])
@login_required
def api_reserve_status(rid):
    if session.get("role") not in ("主席", "副主席"):
        return jsonify({"ok": False, "error": "仅主席/副主席可审批"}), 403
    d = request.get_json(force=True, silent=True) or {}
    st = (d.get("status") or "").strip()
    if st not in ("待确认", "已通过", "已驳回"):
        return jsonify({"ok": False, "error": "状态非法"}), 400
    cx = db(DB_RES)
    cx.execute("UPDATE resv SET status=?, updated=? WHERE id=?",
               (st, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), rid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

# =========== 活动通报子系统（独立库） ===========
@app.route("/work/bulletin")
@login_required
def bulletin_page():
    return send_file(os.path.join(BASE, "bulletin.html"))

@app.route("/api/bulletin")
@login_required
def api_bulletin():
    cx = db(DB_BUL)
    rows = cx.execute("SELECT * FROM bul ORDER BY pinned DESC, id DESC").fetchall()
    cx.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/bulletin", methods=["POST"])
@chair_required
def api_bulletin_save():
    d = request.get_json(force=True, silent=True) or {}
    title = (d.get("title") or "").strip()
    body = (d.get("body") or "").strip()
    level = (d.get("level") or "普通").strip()
    if level not in ("普通", "重要", "紧急"):
        level = "普通"
    if not (title and body):
        return jsonify({"ok": False, "error": "请填写标题与内容"}), 400
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx = db(DB_BUL)
    cols = ["author", "author_xh", "title", "body", "level", "pinned", "created"]
    vals = [session["name"], session["uid"], title, body, level,
            1 if d.get("pinned") else 0, now]
    cx.execute("INSERT INTO bul (" + ",".join(cols) + ") VALUES (" + ",".join("?" * len(cols)) + ")", vals)
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/bulletin/<int:bid>", methods=["DELETE"])
@login_required
def api_bulletin_del(bid):
    cx = db(DB_BUL)
    row = cx.execute("SELECT author_xh FROM bul WHERE id=?", (bid,)).fetchone()
    if row and (row["author_xh"] == session["uid"] or session.get("role") in ("主席", "副主席")):
        cx.execute("DELETE FROM bul WHERE id=?", (bid,)); cx.commit()
    cx.close()
    return jsonify({"ok": True})

# ===== 星海艺术团 · 暗色落地页（React 构建产物，替换旧首页）=====
LANDING = os.path.join(BASE, "portfolio-landing", "landing")

@app.route("/")
def index():
    # 优先返回新的 React 落地页；不存在则回退旧首页（安全兜底）
    landing_html = os.path.join(LANDING, "index.html")
    if os.path.isfile(landing_html):
        return send_file(landing_html)
    return send_file(os.path.join(BASE, "index.html"))

@app.route("/assets/<path:filename>")
def landing_assets(filename):
    return send_file(os.path.join(LANDING, "assets", filename))

# 图片复用线上已有的 static/uploads（不重复存储）
@app.route("/showcase/<path:filename>")
def landing_showcase(filename):
    return send_file(os.path.join(STATIC, "uploads", "showcase", filename))

@app.route("/u20260717171010_f7be1541.png")
def landing_qr():
    return send_file(os.path.join(STATIC, "uploads", "u20260717171010_f7be1541.png"))


# ============ 注册申请 ============
def pwd_strong(pwd):
    return len(pwd) >= 6 and any(c.isdigit() for c in pwd) and any(c.isalpha() for c in pwd)

@app.route("/register")
def register_page():
    return send_file(os.path.join(BASE, "register.html"))

@app.route("/work/approval")
@login_required
@chair_required
def approval_page():
    return send_file(os.path.join(BASE, "approval.html"))

@app.route("/api/register", methods=["POST"])
def api_register_submit():
    d = request.get_json(force=True, silent=True) or {}
    xh = (d.get("xh") or "").strip()
    pwd = (d.get("pwd") or "").strip()
    name = (d.get("name") or "").strip()
    campus = (d.get("campus") or "").strip()
    if not (xh and pwd and name and campus):
        return jsonify({"ok": False, "error": "请填写学号、密码、姓名与校区"}), 400
    if not (xh.isdigit() and len(xh) == 9):
        return jsonify({"ok": False, "error": "学号须为 9 位数字"}), 400
    if not pwd_strong(pwd):
        return jsonify({"ok": False, "error": "密码须至少 6 位且包含字母和数字"}), 400
    cx = db(DB_APP)
    exist = cx.execute("SELECT 1 FROM applies WHERE xh=?", (xh,)).fetchone()
    if exist:
        cx.close()
        return jsonify({"ok": False, "error": "该学号已申请"}), 400
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx.execute("INSERT INTO applies (xh,pwd,name,campus,status,created,updated) VALUES (?,?,?,?,?,?,?)",
               (xh, hx(pwd), name, campus, "待审", now, now))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/register")
@login_required
@chair_required
def api_register_list():
    cx = db(DB_APP)
    rows = cx.execute("SELECT * FROM applies ORDER BY id DESC").fetchall()
    cx.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/register/<int:aid>/approve", methods=["POST"])
@login_required
@chair_required
def api_register_approve(aid):
    cx = db(DB_APP)
    row = cx.execute("SELECT * FROM applies WHERE id=?", (aid,)).fetchone()
    if not row:
        cx.close()
        return jsonify({"ok": False, "error": "申请不存在"}), 404
    cx2 = db(DB_USR)
    cx2.execute("INSERT OR IGNORE INTO users (xh,name,role,pwd,campus,status) VALUES (?,?,?,?,?,?)",
                (row["xh"], row["name"], "成员", row["pwd"], row["campus"], "在团"))
    cx2.commit(); cx2.close()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx.execute("UPDATE applies SET status='已通过', updated=? WHERE id=?", (now, aid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/register/<int:aid>/reject", methods=["POST"])
@login_required
@chair_required
def api_register_reject(aid):
    cx = db(DB_APP)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx.execute("UPDATE applies SET status='已驳回', updated=? WHERE id=?", (now, aid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

# ============ 报销导出 ============
@app.route("/api/reimburse/<int:rid>/export")
@login_required
def api_reimburse_export(rid):
    from docx import Document
    cx = db(DB_REI)
    row = cx.execute("SELECT * FROM rei WHERE id=?", (rid,)).fetchone()
    cx.close()
    if not row:
        return jsonify({"ok": False, "error": "记录不存在"}), 404
    doc = Document()
    doc.add_heading("星海艺术团报销单", 0)
    doc.add_paragraph("申请人：%s（%s）" % (row["name"], row["xh"]))
    doc.add_paragraph("事项：%s" % row["title"])
    doc.add_paragraph("类别：%s" % row["category"])
    doc.add_paragraph("金额：%s 元" % row["amount"])
    doc.add_paragraph("备注：%s" % (row.get("note") or ""))
    doc.add_paragraph("状态：%s" % row["status"])
    doc.add_paragraph("提交时间：%s" % row["created"])
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="%s_%s_%s.docx" % (row["name"], row["title"], rid),
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ============ 资产管理子系统 ============
@app.route("/work/assets")
@login_required
def assets_page():
    return send_file(os.path.join(BASE, "assets.html"))

@app.route("/api/assets")
@login_required
def api_assets():
    cx = db(DB_AST)
    rows = cx.execute("SELECT * FROM assets ORDER BY id DESC").fetchall()
    cx.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/assets", methods=["POST"])
@login_required
def api_assets_save():
    d = request.get_json(force=True, silent=True) or {}
    name = (d.get("name") or "").strip()
    category = (d.get("category") or "").strip()
    if not (name and category):
        return jsonify({"ok": False, "error": "请填写名称与类别"}), 400
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx = db(DB_AST)
    cx.execute("INSERT INTO assets (name,category,qty,location,person,status,created,updated) VALUES (?,?,?,?,?,?,?,?)",
               (name, category, int(d.get("qty") or 1),
                (d.get("location") or "").strip(), (d.get("person") or "").strip(),
                "在库", now, now))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/assets/<int:aid>", methods=["PUT"])
@login_required
def api_assets_update(aid):
    d = request.get_json(force=True, silent=True) or {}
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx = db(DB_AST)
    cx.execute("UPDATE assets SET name=?,category=?,qty=?,location=?,person=?,updated=? WHERE id=?",
               ((d.get("name") or "").strip(), (d.get("category") or "").strip(),
                int(d.get("qty") or 1), (d.get("location") or "").strip(),
                (d.get("person") or "").strip(), now, aid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/assets/<int:aid>", methods=["DELETE"])
@login_required
@chair_required
def api_assets_del(aid):
    cx = db(DB_AST)
    cx.execute("DELETE FROM assets WHERE id=?", (aid,))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/assets/<int:aid>/borrow", methods=["POST"])
@login_required
def api_assets_borrow(aid):
    d = request.get_json(force=True, silent=True) or {}
    cx = db(DB_AST)
    asset = cx.execute("SELECT * FROM assets WHERE id=?", (aid,)).fetchone()
    if not asset:
        cx.close()
        return jsonify({"ok": False, "error": "资产不存在"}), 404
    qty = int(d.get("qty") or 1)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx.execute("INSERT INTO borrows (asset_id,asset_name,xh,name,qty,note,borrow_date,status,created) VALUES (?,?,?,?,?,?,?,?,?)",
               (aid, asset["name"], session["uid"], session["name"], qty,
                (d.get("note") or "").strip(), now, "借出", now))
    cx.execute("UPDATE assets SET status='借出', updated=? WHERE id=?", (now, aid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/assets/<int:aid>/return", methods=["POST"])
@login_required
def api_assets_return(aid):
    cx = db(DB_AST)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx.execute("UPDATE borrows SET status='已还', return_date=? WHERE asset_id=? AND status='借出'",
               (now, aid))
    cx.execute("UPDATE assets SET status='在库', updated=? WHERE id=?", (now, aid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/assets/<int:aid>/repair", methods=["POST"])
@login_required
def api_assets_repair(aid):
    cx = db(DB_AST)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx.execute("UPDATE assets SET status='维修中', updated=? WHERE id=?", (now, aid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/assets/<int:aid>/scrap", methods=["POST"])
@login_required
@chair_required
def api_assets_scrap(aid):
    cx = db(DB_AST)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx.execute("UPDATE assets SET status='已报废', updated=? WHERE id=?", (now, aid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/borrows")
@login_required
def api_borrows():
    cx = db(DB_AST)
    rows = cx.execute("SELECT * FROM borrows ORDER BY id DESC").fetchall()
    cx.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/procurements")
@login_required
def api_procurements():
    cx = db(DB_AST)
    rows = cx.execute("SELECT * FROM procurements ORDER BY id DESC").fetchall()
    cx.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/procurements", methods=["POST"])
@login_required
def api_procurements_save():
    d = request.get_json(force=True, silent=True) or {}
    name = (d.get("name") or "").strip()
    category = (d.get("category") or "").strip()
    if not (name and category):
        return jsonify({"ok": False, "error": "请填写名称与类别"}), 400
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx = db(DB_AST)
    cx.execute("INSERT INTO procurements (name,category,qty,reason,xh,pname,status,created,updated) VALUES (?,?,?,?,?,?,?,?,?)",
               (name, category, int(d.get("qty") or 1),
                (d.get("reason") or "").strip(), session["uid"], session["name"],
                "待审", now, now))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/procurements/<int:pid>/approve", methods=["POST"])
@login_required
@chair_required
def api_procurements_approve(pid):
    cx = db(DB_AST)
    row = cx.execute("SELECT * FROM procurements WHERE id=?", (pid,)).fetchone()
    if not row:
        cx.close()
        return jsonify({"ok": False, "error": "采购申请不存在"}), 404
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx.execute("INSERT INTO assets (name,category,qty,location,person,status,created,updated) VALUES (?,?,?,?,?,?,?,?)",
               (row["name"], row["category"], row["qty"], "", row["pname"], "在库", now, now))
    cx.execute("UPDATE procurements SET status='已通过', updated=? WHERE id=?", (now, pid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/procurements/<int:pid>/reject", methods=["POST"])
@login_required
@chair_required
def api_procurements_reject(pid):
    cx = db(DB_AST)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cx.execute("UPDATE procurements SET status='已驳回', updated=? WHERE id=?", (now, pid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

# ============ 知识库 ============
@app.route("/api/kb")
def api_kb():
    if os.path.exists(KB_FILE):
        with open(KB_FILE, encoding="utf-8") as f:
            return jsonify(json.load(f))
    return jsonify([])

@app.route("/api/kb", methods=["POST"])
def api_kb_save():
    if request.args.get("key") != ADMIN_KEY and (request.get_json(silent=True) or {}).get("key") != ADMIN_KEY:
        return jsonify({"ok": False, "error": "无权限"}), 403
    d = request.get_json(force=True, silent=True) or {}
    keywords = d.get("keywords", [])
    answer = (d.get("answer") or "").strip()
    if not (keywords and answer):
        return jsonify({"ok": False, "error": "请填写关键词和答案"}), 400
    data = json.load(open(KB_FILE, encoding="utf-8")) if os.path.exists(KB_FILE) else []
    nid = max([it.get("id", 0) for it in data], default=0) + 1
    data.append({"id": nid, "keywords": keywords, "answer": answer})
    with open(KB_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return jsonify({"ok": True, "id": nid})

@app.route("/api/kb/<int:kid>", methods=["PUT"])
def api_kb_update(kid):
    if request.args.get("key") != ADMIN_KEY and (request.get_json(silent=True) or {}).get("key") != ADMIN_KEY:
        return jsonify({"ok": False, "error": "无权限"}), 403
    d = request.get_json(force=True, silent=True) or {}
    data = json.load(open(KB_FILE, encoding="utf-8")) if os.path.exists(KB_FILE) else []
    for it in data:
        if it.get("id") == kid:
            if d.get("keywords"):
                it["keywords"] = d["keywords"]
            if d.get("answer"):
                it["answer"] = d["answer"].strip()
            with open(KB_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "条目不存在"}), 404

@app.route("/api/kb/<int:kid>", methods=["DELETE"])
def api_kb_del(kid):
    if request.args.get("key") != ADMIN_KEY and (request.get_json(silent=True) or {}).get("key") != ADMIN_KEY:
        return jsonify({"ok": False, "error": "无权限"}), 403
    data = json.load(open(KB_FILE, encoding="utf-8")) if os.path.exists(KB_FILE) else []
    data = [it for it in data if it.get("id") != kid]
    with open(KB_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return jsonify({"ok": True})

# ============ 数据统计 ============
@app.route("/api/stats")
@login_required
def api_stats():
    cx_mem = db(DB_MEM)
    mem = cx_mem.execute("SELECT COUNT(*) AS c FROM members WHERE status='在团'").fetchone()["c"]
    cx_mem.close()
    cx_rei = db(DB_REI)
    rei = cx_rei.execute("SELECT COUNT(*) AS c FROM rei WHERE status='待审'").fetchone()["c"]
    cx_rei.close()
    cx_app = db(DB_APP)
    apply = cx_app.execute("SELECT COUNT(*) AS c FROM applies WHERE status='待审'").fetchone()["c"]
    cx_app.close()
    cx_ast = db(DB_AST)
    ast = cx_ast.execute("SELECT COUNT(*) AS c FROM assets").fetchone()["c"]
    proc = cx_ast.execute("SELECT COUNT(*) AS c FROM procurements WHERE status='待审'").fetchone()["c"]
    cx_ast.close()
    return jsonify({"ok":True, "members": mem, "assets": ast, "reimburse": rei, "apply": apply, "procure": proc})

# ============ 成员管理扩展 ============
@app.route("/api/signup/<int:rid>/status", methods=["POST"])
@login_required
def api_signup_status(rid):
    d = request.get_json(force=True, silent=True) or {}
    st = (d.get("status") or "").strip()
    if st not in ("通过", "淘汰"):
        return jsonify({"ok": False, "error": "状态须为通过或淘汰"}), 400
    cx = db(DB_REG)
    try:
        cx.execute("ALTER TABLE registrations ADD COLUMN interview_status TEXT DEFAULT '待面试'")
    except sqlite3.OperationalError:
        pass
    cx.execute("UPDATE registrations SET interview_status=? WHERE id=?", (st, rid))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/signup/<int:rid>/enroll", methods=["POST"])
@chair_required
def api_signup_enroll(rid):
    cx = db(DB_REG)
    row = cx.execute("SELECT * FROM registrations WHERE id=?", (rid,)).fetchone()
    if not row:
        cx.close()
        return jsonify({"ok": False, "error": "报名记录不存在"}), 404
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    temp_xh = "8" + str(rid).zfill(8)[:8]
    cx2 = db(DB_USR)
    cx2.execute("INSERT OR IGNORE INTO users (xh,name,role,pwd,campus,status) VALUES (?,?,?,?,?,?)",
                (temp_xh, row["name"], "成员", hx("123456"), (row.get("campus") or ""), "在团"))
    cx2.commit(); cx2.close()
    cx3 = db(DB_MEM)
    cx3.execute("INSERT OR IGNORE INTO members (xh,name,gender,campus,college,major,phone,wechat,email,skill,position,dept,join_date,status,updated) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (temp_xh, row["name"], (row.get("gender") or ""), (row.get("campus") or ""),
                 (row.get("college") or ""), (row.get("major") or ""), (row.get("phone") or ""),
                 (row.get("wechat") or ""), (row.get("email") or ""),
                 (row.get("skill") or ""), "成员", (row.get("target") or ""),
                 now, "在团", now))
    cx3.commit(); cx3.close()
    cx.close()
    return jsonify({"ok": True, "xh": temp_xh})

@app.route("/api/signup/<int:rid>/archive", methods=["POST"])
@login_required
def api_signup_archive(rid):
    cx = db(DB_REG)
    row = cx.execute("SELECT * FROM registrations WHERE id=?", (rid,)).fetchone()
    if not row:
        cx.close()
        return jsonify({"ok": False, "error": "报名记录不存在"}), 404
    cx.execute("""CREATE TABLE IF NOT EXISTS archive_regs(
        id INTEGER, time TEXT, target TEXT, name TEXT, gender TEXT, birth TEXT,
        campus TEXT, college TEXT, major TEXT, phone TEXT, wechat TEXT, email TEXT,
        skill TEXT, motive TEXT, adjust INTEGER)""")
    cols = ["id","time","target","name","gender","birth","campus","college","major","phone","wechat","email","skill","motive","adjust"]
    vals = [row[c] for c in cols]
    cx.execute("INSERT INTO archive_regs (" + ",".join(cols) + ") VALUES (" + ",".join("?" * len(cols)) + ")", vals)
    cx.execute("DELETE FROM registrations WHERE id=?", (rid,))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/members/<xh>/leave", methods=["POST"])
@login_required
@chair_required
def api_members_leave(xh):
    cx = db(DB_MEM)
    row = cx.execute("SELECT * FROM members WHERE xh=?", (xh,)).fetchone()
    if not row:
        cx.close()
        return jsonify({"ok": False, "error": "成员不存在"}), 404
    cx.execute("""CREATE TABLE IF NOT EXISTS alumni(
        xh TEXT, name TEXT, gender TEXT, campus TEXT, college TEXT, major TEXT,
        phone TEXT, wechat TEXT, email TEXT, role TEXT, dept TEXT,
        join_date TEXT, grade TEXT, position TEXT, leave_date TEXT, note TEXT)""")
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cols = ["xh","name","gender","campus","college","major","phone","wechat","email","role","dept","join_date","grade","position"]
    vals = [row.get(c, "") for c in cols]
    cx.execute("INSERT INTO alumni (" + ",".join(cols) + ",leave_date,note) VALUES (" + ",".join("?" * (len(cols) + 2)) + ")",
               vals + [now, "离团"])
    cx.execute("DELETE FROM members WHERE xh=?", (xh,))
    cx.commit(); cx.close()
    return jsonify({"ok": True})

@app.route("/api/alumni")
@login_required
def api_alumni():
    cx = db(DB_MEM)
    cx.execute("""CREATE TABLE IF NOT EXISTS alumni(
        xh TEXT, name TEXT, gender TEXT, campus TEXT, college TEXT, major TEXT,
        phone TEXT, wechat TEXT, email TEXT, role TEXT, dept TEXT,
        join_date TEXT, grade TEXT, position TEXT, leave_date TEXT, note TEXT)""")
    rows = cx.execute("SELECT * FROM alumni ORDER BY leave_date DESC").fetchall()
    cx.close()
    return jsonify([dict(r) for r in rows])

# ============ 启动 ============
def init_site():
    cx = db(DB_SITE)
    cx.execute("""CREATE TABLE IF NOT EXISTS site_audit(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ts TEXT, uid TEXT, name TEXT, ip TEXT,
        action TEXT, target TEXT, detail TEXT)""")
    cx.execute("""CREATE TABLE IF NOT EXISTS site_guard(
        k TEXT PRIMARY KEY, fails INTEGER DEFAULT 0,
        first_fail REAL, locked_until REAL)""")
    cx.commit(); cx.close()
init_site()

init_reg(); init_mem(); init_usr(); init_rei(); init_res(); init_bul(); export_csv()

# ============ 官网内容编辑（二次口令体系，施工总案 A6）============
SITE_CONTENT = os.path.join(BASE, "site_content.json")  # 二期运行时内容（预留）

@app.route("/work/site")
@login_required
def work_site_page():
    return send_file(os.path.join(BASE, "site.html"))

@app.route("/api/site/status")
@login_required
def api_site_status():
    ok = _site_unlocked()
    locked_until, fails = _guard_state()
    return jsonify({
        "ok": True,
        "unlocked": ok,
        "configured": bool(SITE_EDIT_KEY),
        "expires_in": int(max(0, float(session.get("site_edit_until", 0)) - time.time())) if ok else 0,
        "locked_for": int(max(0, locked_until - time.time())),
        "fails": fails,
        "max_fail": SITE_MAX_FAIL,
        "role": session.get("role", ""),
        "can_edit": session.get("role") in ("主席", "副主席"),
    })

@app.route("/api/site/unlock", methods=["POST"])
@login_required
def api_site_unlock():
    if session.get("role") not in ("主席", "副主席"):
        return jsonify({"ok": False, "error": "仅主席/副主席可进入官网内容编辑"}), 403
    if not SITE_EDIT_KEY:
        return jsonify({"ok": False, "error": "编辑口令未配置（config.json 缺 SITE_EDIT_KEY）"}), 503
    locked_until, _ = _guard_state()
    if locked_until and time.time() < locked_until:
        mins = int(locked_until - time.time()) // 60 + 1
        return jsonify({"ok": False, "error": "尝试过于频繁，请 %d 分钟后再试" % mins}), 429
    data = request.get_json(silent=True) or {}
    pwd = str(data.get("pwd") or "")
    # 恒定时间比较，防时序侧信道
    if not pwd or not hmac.compare_digest(pwd, SITE_EDIT_KEY):
        n, locked = _guard_fail()
        if locked:
            return jsonify({"ok": False,
                            "error": "连续 %d 次错误，已锁定 %d 分钟" % (SITE_MAX_FAIL, SITE_LOCK_SEC // 60)}), 429
        return jsonify({"ok": False,
                        "error": "口令错误，还可尝试 %d 次" % (SITE_MAX_FAIL - n)}), 403
    _guard_reset()
    session["site_edit_until"] = time.time() + SITE_EDIT_TTL
    session["site_edit_uid"] = session.get("uid")
    site_audit("UNLOCK", "", "二次验证通过")
    return jsonify({"ok": True, "expires_in": SITE_EDIT_TTL})

@app.route("/api/site/lock", methods=["POST"])
@login_required
def api_site_lock():
    session.pop("site_edit_until", None)
    session.pop("site_edit_uid", None)
    site_audit("LOCK_MANUAL", "", "主动上锁")
    return jsonify({"ok": True})

# 字段白名单：content.json 顶层可编辑键（防任意键注入污染）
SITE_CONTENT_KEYS = {"hero_welcome", "b0", "b1", "b2", "b3", "t_wei", "t_dai",
                     "footer", "consult_qr", "team_qr", "dept_desc", "carousel", "org"}
SITE_CONTENT_TYPES = {"team_qr": dict, "dept_desc": dict, "carousel": list, "org": dict}

@app.route("/api/site/content")
@login_required
@site_edit_required
def api_site_content_get():
    return jsonify({"ok": True, "data": load_content()})

@app.route("/api/site/content", methods=["POST"])
@login_required
@site_edit_required
def api_site_content_set():
    data = request.get_json(force=True, silent=True) or {}
    for k in list(data.keys()):
        if k.startswith("_"):
            data.pop(k); continue
        if k not in SITE_CONTENT_KEYS:
            return jsonify({"ok": False, "error": "未知字段: " + k}), 400
        if k in SITE_CONTENT_TYPES and not isinstance(data[k], SITE_CONTENT_TYPES[k]):
            return jsonify({"ok": False, "error": "字段类型错误: " + k}), 400
    cur = load_content()
    changes = []
    for k, v in data.items():
        old = cur.get(k)
        if old != v:
            changes.append({"field": k, "old": _brief(old), "new": _brief(v)})
    if not changes:
        return jsonify({"ok": True, "changed": 0})
    cur.update(data)
    save_content_atomic(CONTENT, cur)
    site_audit("SAVE", ",".join(c["field"] for c in changes)[:200],
               json.dumps(changes, ensure_ascii=False))
    return jsonify({"ok": True, "changed": len(changes)})

@app.route("/api/site/upload", methods=["POST"])
@login_required
@site_edit_required
def api_site_upload():
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "未收到文件"}), 400
    parts = (f.filename or "img.png").rsplit(".", 1)
    ext = (parts[-1] if len(parts) > 1 else "png").lower()
    if ext not in ("png", "jpg", "jpeg", "gif", "webp"):
        return jsonify({"ok": False, "error": "仅支持 png/jpg/jpeg/gif/webp"}), 400
    raw = f.read()
    if len(raw) > 5 * 1024 * 1024:
        return jsonify({"ok": False, "error": "图片不能超过 5MB"}), 400
    if not sniff_image(raw):
        return jsonify({"ok": False, "error": "文件内容不是有效图片"}), 400
    f.stream = io.BytesIO(raw)

    slot = request.form.get("slot")  # 首页槽位覆盖（保留文件名=免 build 生效）
    if slot:
        if not re.fullmatch(r"sc\d{2}", slot):
            return jsonify({"ok": False, "error": "槽位名非法（应为 sc01~sc99）"}), 400
        dst_dir = os.path.join(STATIC, "uploads", "showcase")
        os.makedirs(dst_dir, exist_ok=True)
        dst = os.path.join(dst_dir, slot + ".jpg")
        if os.path.abspath(dst) != os.path.join(os.path.abspath(dst_dir), slot + ".jpg"):
            return jsonify({"ok": False, "error": "路径非法"}), 400
        if os.path.exists(dst):  # 覆盖前自动备份原图
            sb = os.path.join(BASE, "data", "showcase_backups")
            os.makedirs(sb, exist_ok=True)
            shutil.copy2(dst, os.path.join(sb, "%s-%s.jpg" % (slot, datetime.now().strftime("%Y%m%d-%H%M%S"))))
        f.save(dst)
        site_audit("UPLOAD_SLOT", slot, "覆盖首页图片槽位")
        return jsonify({"ok": True, "url": "/showcase/" + slot + ".jpg"})

    safe = "u" + datetime.now().strftime("%Y%m%d%H%M%S") + "_" + os.urandom(4).hex() + "." + ext
    f.save(os.path.join(STATIC, "uploads", safe))
    site_audit("UPLOAD", safe, "")
    return jsonify({"ok": True, "url": "/static/uploads/" + safe})

@app.route("/api/site/backups")
@login_required
@site_edit_required
def api_site_backups():
    os.makedirs(SITE_BACKUP_DIR, exist_ok=True)
    fs = sorted([f for f in os.listdir(SITE_BACKUP_DIR) if f.startswith("content-")], reverse=True)
    return jsonify({"ok": True, "files": fs[:SITE_BACKUP_KEEP]})

@app.route("/api/site/rollback", methods=["POST"])
@login_required
@site_edit_required
def api_site_rollback():
    data = request.get_json(silent=True) or {}
    fn = str(data.get("file") or "")
    if not fn.startswith("content-") or "/" in fn or "\\" in fn or ".." in fn:
        return jsonify({"ok": False, "error": "文件名非法"}), 400
    src = os.path.join(SITE_BACKUP_DIR, os.path.basename(fn))
    if not os.path.isfile(src):
        return jsonify({"ok": False, "error": "备份不存在"}), 404
    if _backup_content(CONTENT) is None:  # 回滚前先备份当前状态
        return jsonify({"ok": False, "error": "当前内容文件缺失，拒绝盲回滚"}), 500
    shutil.copy2(src, CONTENT)
    _pub_cache["ts"] = 0.0
    site_audit("ROLLBACK", os.path.basename(fn), "回滚到历史版本")
    return jsonify({"ok": True})

@app.route("/api/site/audit")
@login_required
@chair_required
def api_site_audit():
    cx = db(DB_SITE)
    rows = cx.execute("SELECT * FROM site_audit ORDER BY id DESC LIMIT 200").fetchall()
    cx.close()
    return jsonify({"ok": True, "rows": [dict(r) for r in rows]})

# ===== 错误页（施工总案 A5 / 413 归 A1）=====
@app.errorhandler(413)
def e413(e):
    return jsonify({"ok": False, "error": "上传内容超过 10MB 上限"}), 413

@app.errorhandler(404)
def e404(e):
    if request.path.startswith("/api/"):
        return jsonify({"ok": False, "error": "接口不存在"}), 404
    try:
        return send_file(os.path.join(BASE, "404.html")), 404
    except Exception:
        return "404 页面不存在", 404

@app.errorhandler(500)
def e500(e):
    app.logger.exception(e)
    if request.path.startswith("/api/"):
        return jsonify({"ok": False, "error": "服务器内部错误"}), 500
    try:
        return send_file(os.path.join(BASE, "500.html")), 500
    except Exception:
        return "500 服务器内部错误", 500

@app.errorhandler(FileNotFoundError)
def e_fnf(e):
    # send_file 对缺失文件抛 FileNotFoundError（非 HTTPException），按 404 语义处理
    return e404(e)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    app.run(host="0.0.0.0", port=port, debug=False)
